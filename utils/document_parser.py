"""Document parser for extracting text from various file formats."""

import os
from typing import Tuple
from PyPDF2 import PdfReader
from docx import Document


class DocumentParser:
    """Extract text from TXT, PDF, and DOCX files."""
    
    SUPPORTED_FORMATS = {'.txt', '.pdf', '.docx', '.doc'}
    
    @staticmethod
    def parse_file(file_path: str) -> Tuple[str, str]:
        """Parse a document file and extract text.
        
        Args:
            file_path: Path to the document file.
            
        Returns:
            Tuple of (extracted_text, file_format)
            
        Raises:
            ValueError: If file format is not supported.
            Exception: If parsing fails.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext not in DocumentParser.SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported file format: {file_ext}. "
                f"Supported formats: {', '.join(DocumentParser.SUPPORTED_FORMATS)}"
            )
        
        if file_ext == '.txt':
            return DocumentParser._parse_txt(file_path), 'txt'
        elif file_ext == '.pdf':
            return DocumentParser._parse_pdf(file_path), 'pdf'
        elif file_ext in ['.docx', '.doc']:
            return DocumentParser._parse_docx(file_path), 'docx'
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    @staticmethod
    def parse_bytes(file_bytes: bytes, filename: str) -> Tuple[str, str]:
        """Parse a document from bytes.
        
        Args:
            file_bytes: File content as bytes.
            filename: Original filename (to determine format).
            
        Returns:
            Tuple of (extracted_text, file_format)
        """
        import tempfile
        
        file_ext = os.path.splitext(filename)[1].lower()
        
        if file_ext not in DocumentParser.SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported file format: {file_ext}. "
                f"Supported formats: {', '.join(DocumentParser.SUPPORTED_FORMATS)}"
            )
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp_file:
            tmp_file.write(file_bytes)
            tmp_path = tmp_file.name
        
        try:
            text, format_type = DocumentParser.parse_file(tmp_path)
            return text, format_type
        finally:
            # Clean up temporary file
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    
    @staticmethod
    def _parse_txt(file_path: str) -> str:
        """Parse text file.
        
        Args:
            file_path: Path to TXT file.
            
        Returns:
            Extracted text.
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """Parse PDF file.
        
        Args:
            file_path: Path to PDF file.
            
        Returns:
            Extracted text from all pages.
        """
        try:
            reader = PdfReader(file_path)
            text_parts = []
            
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            if not text_parts:
                raise ValueError("No text could be extracted from PDF")
            
            return '\n'.join(text_parts)
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
    
    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """Parse DOCX file.
        
        Args:
            file_path: Path to DOCX file.
            
        Returns:
            Extracted text from all paragraphs.
        """
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            if not text_parts:
                raise ValueError("No text could be extracted from DOCX")
            
            return '\n'.join(text_parts)
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")
    
    @staticmethod
    def get_file_info(file_path: str) -> dict:
        """Get information about a file.
        
        Args:
            file_path: Path to file.
            
        Returns:
            Dictionary with file information.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        stat = os.stat(file_path)
        file_ext = os.path.splitext(file_path)[1].lower()
        
        return {
            'filename': os.path.basename(file_path),
            'size_bytes': stat.st_size,
            'size_mb': round(stat.st_size / (1024 * 1024), 2),
            'format': file_ext,
            'supported': file_ext in DocumentParser.SUPPORTED_FORMATS
        }
